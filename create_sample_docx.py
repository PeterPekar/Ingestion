from docx import Document

# Create a new Document
document = Document()

# Add a heading
document.add_heading('This is a sample DOCX file.', level=1)

# Add a paragraph
document.add_paragraph('It has some text.')

# Add a table
table = document.add_table(rows=1, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Name'
hdr_cells[1].text = 'Age'
row_cells = table.add_row().cells
row_cells[0].text = 'Alice'
row_cells[1].text = '30'
row_cells = table.add_row().cells
row_cells[0].text = 'Bob'
row_cells[1].text = '25'

# Save the document
document.save('documents/sample.docx')
